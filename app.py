import streamlit as st
from datetime import datetime
from docx import Document

st.set_page_config(
    page_title="BookForge AI",
    page_icon="📚",
    layout="wide"
)

st.title("📚 BookForge AI")
st.caption("Turn one idea into a complete book blueprint.")

st.header("Book Information")

idea = st.text_area(
    "Describe your book idea",
    placeholder="Example: A detective discovers dreams predict future crimes..."
)

genre = st.selectbox(
    "Genre",
    [
        "Fantasy",
        "Thriller",
        "Sci-Fi",
        "Romance",
        "Mystery",
        "Horror",
        "Biography",
        "Self Help",
        "Business",
        "Education"
    ]
)

audience = st.selectbox(
    "Target Audience",
    [
        "Children",
        "Teenagers",
        "Young Adults",
        "Adults"
    ]
)

chapters = st.slider(
    "Number of Chapters",
    5,
    30,
    10
)

pages = st.slider(
    "Approximate Pages",
    50,
    500,
    150
)

tone = st.selectbox(
    "Writing Style",
    [
        "Professional",
        "Friendly",
        "Dark",
        "Inspirational",
        "Humorous"
    ]
)

if st.button("🚀 Generate Book Blueprint"):

    title = "Untitled Book"

    if idea:
        title = idea.split(" ")[0].capitalize() + " Chronicles"

    toc = []

    for i in range(1, chapters + 1):
        toc.append(f"Chapter {i}")

    st.success("Blueprint Generated!")

    st.subheader("📖 Book Title")
    st.write(title)

    st.subheader("📋 Table of Contents")

    for chapter in toc:
        st.write("•", chapter)

    st.subheader("📝 Summary")

    st.write(idea)

    document = Document()

    document.add_heading(title, level=1)

    document.add_heading("Book Summary", level=2)
    document.add_paragraph(idea)

    document.add_heading("Table of Contents", level=2)

    for chapter in toc:
        document.add_paragraph(chapter)

    filename = f"{title}.docx"

    document.save(filename)

    with open(filename, "rb") as file:
        st.download_button(
            "📥 Download DOCX",
            file,
            file_name=filename
        )
